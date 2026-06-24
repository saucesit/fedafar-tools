#!/usr/bin/env python3
"""
sc_pliego_scraper.py — Descarga y parsea pliegos de SaltaCompra con Playwright.
Corre localmente (no en Render). Requiere: playwright install chromium

Busca licitaciones SaltaCompra sin items_detalle, descarga el pliego PDF
de cada una, extrae los renglones con Claude y los guarda en Supabase.
"""

import os, sys, json, re, time, base64
from pathlib import Path

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

from dotenv import load_dotenv
load_dotenv(Path(__file__).parent / '.env')

from playwright.sync_api import sync_playwright
from supabase import create_client
import anthropic

from match_catalogo import cargar_terminos_catalogo, matchear_items

SUPABASE_URL  = os.environ['SUPABASE_URL']
SUPABASE_KEY  = os.environ['SUPABASE_KEY']
ANTHROPIC_KEY = os.environ['ANTHROPIC_API_KEY']
# Service key (secreta): necesaria para subir pliegos a Storage. Si no está,
# el guardado del PDF se saltea y el botón Pliego queda sin archivo.
SUPABASE_SERVICE_KEY = os.environ.get('SUPABASE_SERVICE_KEY', '')
# Login ya no es necesario (listado y pliegos públicos); se dejan por compat.
SC_USER       = os.environ.get('SC_USER', '')
SC_PASS       = os.environ.get('SC_PASS', '')

URL_LISTA = 'https://saltacompra.gob.ar/Compras.aspx?qs=iouVZE0yWCs='
TMP_DIR   = Path(__file__).parent / 'tmp_sc_pliegos'
BUCKET    = 'pliegos'

_CTYPES = {
    '.pdf':  'application/pdf',
    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    '.xls':  'application/vnd.ms-excel',
}

# ── Storage: subir el pliego y devolver URL pública ───────────────────────────

def get_storage_client():
    """Cliente con service key para Storage (la anon key no puede subir)."""
    if not SUPABASE_SERVICE_KEY:
        return None
    return create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)

def asegurar_bucket(sb_admin):
    if sb_admin is None:
        return
    try:
        nombres = [b.name for b in sb_admin.storage.list_buckets()]
        if BUCKET not in nombres:
            sb_admin.storage.create_bucket(BUCKET, options={'public': True})
            print(f'  Bucket "{BUCKET}" creado.')
    except Exception as e:
        print(f'  [storage] No se pudo asegurar el bucket: {e}')

def subir_pliego(sb_admin, lic_id, tmp_path):
    """Sube el pliego (upsert) y devuelve su URL pública, o None."""
    if sb_admin is None:
        return None
    ext  = tmp_path.suffix.lower()
    dest = f'{lic_id}{ext}'
    try:
        sb_admin.storage.from_(BUCKET).upload(
            dest, tmp_path.read_bytes(),
            {'content-type': _CTYPES.get(ext, 'application/octet-stream'), 'upsert': 'true'}
        )
        return sb_admin.storage.from_(BUCKET).get_public_url(dest)
    except Exception as e:
        print(f'    [storage] Error subiendo: {e}')
        return None

# ── Parsear pliego con Claude (PDF / Word / Excel) ────────────────────────────

# ── Parsear pliego con Claude (PDF / Word / Excel) ────────────────────────────

_PROMPT_ITEMS = (
    'Este es el pliego de una licitacion publica argentina. Objeto: {objeto}\n\n'
    'Extraé la tabla de renglones/items que se solicitan comprar. '
    'Para cada item incluí descripcion, cantidad y unidad (si existe). '
    'Respondi SOLO con un JSON array sin markdown:\n'
    '[{{"descripcion":"...","cantidad":"...","unidad":"..."}},...]\n'
    'Si no hay tabla de items, respondi [].'
)

def _claude_items(content_blocks):
    """Devuelve la lista de items, [] si el pliego no tiene tabla, o None si
    falló por rate limit persistente (para reintentar en otra corrida)."""
    client = anthropic.Anthropic(api_key=ANTHROPIC_KEY)
    for intento in range(4):
        try:
            resp = client.messages.create(
                model='claude-haiku-4-5-20251001',
                max_tokens=1500,
                messages=[{'role': 'user', 'content': content_blocks}]
            )
            text = resp.content[0].text.strip()
            text = re.sub(r'^```json?\s*', '', text)
            text = re.sub(r'\s*```$', '', text)
            return json.loads(text)
        except Exception as e:
            msg = str(e)
            if '429' in msg or 'rate_limit' in msg.lower():
                espera = 30 * (intento + 1)
                print(f'    [Claude] Rate limit; esperando {espera}s y reintento...')
                time.sleep(espera)
                continue
            print(f'    [Claude] Error: {e}')
            return []
    print('    [Claude] Rate limit persistente; se reintenta en otra corrida')
    return None

def parsear_pdf(pdf_bytes, objeto):
    pdf_b64 = base64.standard_b64encode(pdf_bytes).decode('utf-8')
    return _claude_items([
        {'type': 'document', 'source': {
            'type': 'base64', 'media_type': 'application/pdf', 'data': pdf_b64
        }},
        {'type': 'text', 'text': _PROMPT_ITEMS.format(objeto=objeto)},
    ])

def _parsear_texto(texto, objeto):
    if not texto.strip():
        return []
    return _claude_items([
        {'type': 'text', 'text': _PROMPT_ITEMS.format(objeto=objeto) +
                                 '\n\nCONTENIDO DEL PLIEGO:\n' + texto[:12000]},
    ])

def parsear_pliego(path, objeto):
    """Despacha según extensión: PDF, Word (.docx) o Excel (.xlsx)."""
    ext = path.suffix.lower()
    try:
        if ext == '.pdf':
            return parsear_pdf(path.read_bytes(), objeto)

        if ext == '.docx':
            from docx import Document
            doc   = Document(str(path))
            partes = [p.text for p in doc.paragraphs if p.text.strip()]
            for t in doc.tables:
                for row in t.rows:
                    celdas = [c.text.strip() for c in row.cells]
                    if any(celdas):
                        partes.append(' | '.join(celdas))
            return _parsear_texto('\n'.join(partes), objeto)

        if ext in ('.xlsx', '.xls'):
            import openpyxl
            wb     = openpyxl.load_workbook(str(path), data_only=True)
            lineas = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    fila = ' | '.join(str(c) for c in row if c is not None)
                    if fila.strip():
                        lineas.append(fila)
            return _parsear_texto('\n'.join(lineas[:400]), objeto)
    except Exception as e:
        print(f'    [parsear_pliego] Error con {ext}: {e}')
        return []

    print(f'    Formato {ext} no soportado')
    return []

# ── Buscar la fila recorriendo las páginas ────────────────────────────────────

def buscar_fila(page, numero, objeto, max_paginas=15):
    """Busca la fila de la licitación recorriendo la paginación del listado.
    Devuelve el locator de la fila, o None si no aparece."""
    page.goto(URL_LISTA, wait_until='networkidle', timeout=30000)
    page.wait_for_timeout(800)
    palabras = ' '.join(objeto.split()[:4])

    for pagina in range(1, max_paginas + 1):
        fila = page.locator(f'tr:has-text("{numero}")').first
        if fila.count() > 0:
            return fila
        if palabras:
            fila = page.locator(f'tr:has-text("{palabras}")').first
            if fila.count() > 0:
                return fila

        # Avanzar a la página siguiente (postback AJAX vía clic en el paginador)
        sig  = pagina + 1
        link = page.locator(f"a[href*=\"'Page${sig}'\"]")
        if link.count() == 0:
            break
        try:
            link.first.click()
            page.wait_for_load_state('networkidle', timeout=30000)
            page.wait_for_timeout(900)
        except Exception:
            break
    return None

# ── Pase al CRM si los items matchean el catálogo ─────────────────────────────

def evaluar_crm(sb, lic_id, items, terminos):
    """Si al menos un item matchea el catálogo, agrega la licitación al CRM.
    Devuelve (matcheados, total)."""
    n_match, terms_match = matchear_items(items, terminos)
    if n_match == 0:
        print(f'    => sin coincidencias con catálogo — queda en Licitaciones')
        return 0, len(items)

    existing = sb.table('licitaciones_crm').select('id') \
                 .eq('licitacion_id', str(lic_id)).execute().data
    if not existing:
        nota = (f'{n_match} de {len(items)} items matchean catálogo: '
                f'{", ".join(terms_match[:8])}')
        sb.table('licitaciones_crm').insert({
            'licitacion_id': str(lic_id), 'estado': 'identificada', 'notas': nota
        }).execute()
    print(f'    => {n_match}/{len(items)} items matchean catálogo → agregada al CRM')
    return n_match, len(items)

# ── Main ──────────────────────────────────────────────────────────────────────

def _sin_items(r):
    return not r.get('items_detalle') or r['items_detalle'] in ('[]', 'null', '')

def _tiene_pdf(r):
    return 'storage/v1/object/public' in (r.get('url') or '')

def run(limite=None):
    sb       = create_client(SUPABASE_URL, SUPABASE_KEY)
    sb_admin = get_storage_client()           # service key, para Storage
    storage_ok = sb_admin is not None
    asegurar_bucket(sb_admin)
    TMP_DIR.mkdir(exist_ok=True)

    if not storage_ok:
        print('[aviso] Sin SUPABASE_SERVICE_KEY: no se guardarán los PDFs (botón Pliego sin archivo).\n')

    rows = sb.table('licitaciones').select('id,numero_proceso,objeto,items_detalle,url') \
             .eq('fuente', 'saltacompra') \
             .in_('clasificacion', ['APLICA', 'REVISAR']).execute().data or []

    # Pendiente = sin items, o (si hay Storage) sin el PDF subido todavía (backfill)
    pendientes = [r for r in rows
                  if _sin_items(r) or (storage_ok and not _tiene_pdf(r))]
    if limite:
        pendientes = pendientes[:limite]
    print(f'Licitaciones SC a procesar: {len(pendientes)}')
    if not pendientes:
        print('Nada que procesar.')
        return

    print('Cargando términos del catálogo para matcheo...')
    terminos = cargar_terminos_catalogo()
    print(f'Términos de catálogo: {len(terminos)}\n')

    with sync_playwright() as p:
        browser  = p.chromium.launch(headless=True)
        context  = browser.new_context(accept_downloads=True)
        page     = context.new_page()

        # No hace falta login: el listado y la descarga de pliegos son públicos.
        print('Abriendo SaltaCompra (sin login)...\n')

        ok = sin_pliego = 0
        for lic in pendientes:
            numero = lic['numero_proceso']
            objeto = lic['objeto']
            print(f'  {numero} — {objeto[:55]}')

            # Buscar la fila recorriendo las páginas del listado
            fila = buscar_fila(page, numero, objeto)
            if fila is None:
                print(f'    No encontrado en la lista (puede haber vencido)')
                # Solo marcar vacío si no tenía items; nunca pisar items existentes
                if _sin_items(lic):
                    sb.table('licitaciones').update({'items_detalle': '[]'}).eq('id', lic['id']).execute()
                sin_pliego += 1
                continue

            btn = fila.locator('a:has-text("Pliego"), a:has-text("Descargar")')
            if btn.count() == 0:
                print(f'    Sin boton de descarga')
                sin_pliego += 1
                continue

            try:
                with page.expect_download(timeout=15000) as dl_info:
                    btn.first.click()
                dl = dl_info.value
                tmp = TMP_DIR / dl.suggested_filename
                dl.save_as(str(tmp))
                size = tmp.stat().st_size
                print(f'    Descargado: {tmp.name} ({size:,} bytes)')

                update = {}

                # Subir el pliego a Storage y guardar su URL pública
                url_pliego = subir_pliego(sb_admin, lic['id'], tmp)
                if url_pliego:
                    update['url'] = url_pliego
                    print(f'    Pliego subido a Storage')

                # Extraer items solo si todavía no los tiene (evita re-llamar a Claude)
                if _sin_items(lic):
                    items = parsear_pliego(tmp, objeto)
                    if items is None:   # rate limit: no pisar, se reintenta otra corrida
                        print('    => items pendientes (rate limit), se reintenta luego')
                        if update:
                            sb.table('licitaciones').update(update).eq('id', lic['id']).execute()
                        tmp.unlink(missing_ok=True)
                        sin_pliego += 1
                        time.sleep(1)
                        continue
                    nombres = [i['descripcion'] for i in items if i.get('descripcion')]
                    update['items_detalle']        = json.dumps(items,   ensure_ascii=False)
                    update['productos_detectados'] = json.dumps(nombres, ensure_ascii=False)
                    print(f'    => {len(items)} items extraídos')
                else:
                    items = json.loads(lic['items_detalle'])

                if update:
                    sb.table('licitaciones').update(update).eq('id', lic['id']).execute()
                if items:
                    evaluar_crm(sb, lic['id'], items, terminos)
                ok += 1
                tmp.unlink(missing_ok=True)

            except Exception as e:
                print(f'    [ERROR] {e}')
                sin_pliego += 1

            time.sleep(1)

        browser.close()

    print(f'\n[OK] Con items: {ok} | Sin pliego: {sin_pliego}')


if __name__ == '__main__':
    limite = int(sys.argv[1]) if len(sys.argv) > 1 and sys.argv[1].isdigit() else None
    run(limite)
